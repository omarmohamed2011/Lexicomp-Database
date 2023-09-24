import re
from docx import Document

def get_maximum(dose):
    ##"(maximum: 60 mg/24 hours) g/L mg/l day mg/"
    regex1 = "(M|m)aximum: (\d)+?(.(\d)+?)? (.{2}|.|.{3})/(.|.{2}|.{3}|.{4}|.{5}) (.{5}|.{4}|.{3}|.{2}|.{1})?"
    regex2 = "(M|m)aximum dose: (\d)+?(.(\d)+?)? (.{2}|.|.{3})/(.|.{2}|.{3}|.{4}|.{5}) (.{5}|.{4}|.{3}|.{2}|.{1})?"
    regex3 = "(M|m)aximum daily dose: (\d)+?(.(\d)+?)? (.{2}|.|.{3})/(.|.{2}|.{3}|.{4}|.{5}) (.{5}|.{4}|.{3}|.{2}|.{1})?"
    regex4 = "(M|m)aximum single dose: (\d)+?(.(\d)+?)? (.{2}|.|.{3})"
    regex5 = "(M|m)aximum total dose: (\d)+?(.(\d)+?)? (.{2}|.|.{3})"

    if re.search(regex1,dose):
        return re.search(regex1,dose)
    if re.search(regex2,dose):
        return re.search(regex2,dose)
    if re.search(regex3,dose):
        return re.search(regex3,dose)
    if re.search(regex4,dose):
        return re.search(regex4,dose)
    if re.search(regex5,dose):
        return re.search(regex5,dose)

def return_maximum(dose):
    try:
        span_range = get_maximum(dose).span()
        return dose[span_range[0]:span_range[1]]
    except:
        return ''

def get_heading_text(paragraphs, heading_name):
    found = False
    text = ""
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading 1') and (heading_name in paragraph.text):
            found = True
            continue
        if found and paragraph.style.name != "Heading 1":
            text += paragraph.text
            text += '\n\n'
        elif found and paragraph.style.name == "Heading 1":
            break
    return text


def get_paragraphs_lst(paragraphs, heading_name):
    found = False
    heading_paragraphs=[]
    for paragraph in paragraphs:
        if paragraph.style.name.startswith('Heading 1') and (heading_name in paragraph.text):
            found = True
            continue
        if found and paragraph.style.name != "Heading 1":
            heading_paragraphs.append(paragraph)
        elif found and paragraph.style.name == "Heading 1":
            break
    return heading_paragraphs


def get_disease_txt(paragraphs):
    list_paragraphs = []
    temp=None
    for p in paragraphs:
        text=p.text
        if len(text) >1:
            if text[-1] != '.':
                temp=text
            else:
                if temp is not None:
                    list_paragraphs.append(temp+text)
                    temp=None
                else:
                    list_paragraphs.append(text)

    return list_paragraphs

def get_title(document):
    return document.paragraphs[0].text


def dosing_splitting(list_pdfs):
    disease_name = [];
    disease_description = [];
    dosage_title = []

    set_startings = ['Initial', 'Dosage adjustment', 'Maximum', 'Maintenance', 'IV', 'Oral', 'Oral, IV', 'Rectal',
                     "Consult drug interactions database for more information", '≤', '>', '≥', '<', 'IM', 'If',
                     'Patients', 'Hb', 'Preoperative', 'SUBQ', 'Consider', 'Consult']

    heading_name = 'Dosing: Adult'

    for pdf_path in list_pdfs[:]:

        document = Document(pdf_path)
        title = get_title(document)

        if title == '':
            title = pdf_path

        dosing_adults = get_paragraphs_lst(document.paragraphs, heading_name)
        list_paragraphs = get_disease_txt(dosing_adults)

        for para in list_paragraphs:
            i = 0
            flag = 0
            if ':' not in para:
                flag = 1

            for start_word in set_startings:
                if para.startswith(start_word):
                    flag = 1
            if i > 0:
                main_case = para.startswith('Note:')
                if (main_case and ((para.split(':')[0].split() == disease_name[-1].split()))) or (
                (disease_name[-1].split() in para)):
                    flag = 1
            if flag != 1:
                dosage_title.append(title)
                disease_name.append(para.split(':')[0])
                disease_description.append(para)
            else:
                disease_description[-1] = disease_description[-1] + '\n' + para
            i += 1

    return dosage_title,disease_name,disease_description

def adult_injection(disease_description):
    Oral_adult, IV_adult, Rectal_adult = [], [], []

    for dose in disease_description:

        Oral_adult.append(''); IV_adult.append(''); Rectal_adult.append('')

        for sub_dose in dose:
            flag = 0

            for elem in ['Oral', 'Oral, IV', 'IV, Oral']:
                if elem in sub_dose:
                    flag = 1
            if flag == 1:
                Oral_adult[-1] += sub_dose

            flag = 0
            for elem in ['Oral, IV', 'IV, Oral', 'IV']:
                if elem in sub_dose:
                    flag = 1
            if flag == 1:
                IV_adult[-1] += sub_dose

            if 'Rectal' in sub_dose:
                Rectal_adult[-1] += sub_dose

    return Oral_adult, IV_adult, Rectal_adult

def injection_type(dosage_title,disease_description):
    Oral_IV = []

    for i in range(len(dosage_title)):
        dis = disease_description[i]
        try:
            iv_data = dis.split("Oral, IV")[1]
        except:
            iv_data = ''
        Oral_IV.append(iv_data)

    IV, Oral, Rectal, IM = [], [], [], []

    for i in range(len(dosage_title)):

        dis = disease_description[i]
        if Oral_IV[i] != '':
            IV.append(Oral_IV[i]);
            Oral.append(Oral_IV[i])
            try:
                rectal_data = dis.split("Rectal:")[1]
                IM_data = dis.split("IM:")[1]
            except:
                rectal_data = '';
                IM_data = ''

            Rectal.append(rectal_data)
            IM.append(IM_data)
        else:
            try:
                iv_data = dis.split("IV:")[1]
            except:
                iv_data = ''
            try:
                oral_data = dis.split("Oral:")[1]
            except:
                oral_data = ''
            try:
                rectal_data = dis.split("Rectal:")[1]
            except:
                rectal_data = ''
            try:
                IM_data = dis.split("IM:")[1]
            except:
                IM_data = ''

            IV.append(iv_data);
            Oral.append(oral_data);
            Rectal.append(rectal_data);
            IM.append(IM_data)

    return IV,Oral,Rectal,IM,Oral_IV